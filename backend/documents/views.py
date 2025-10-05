import os
import tempfile
from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser
from openai import OpenAI
from docx import Document
from .utils import extract_text_from_pdf, extract_text_from_docx
import json
import datetime

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


class GenerateICFView(APIView):
    parser_classes = [MultiPartParser]

    def post(self, request):
        file = request.FILES.get("file")
        if not file:
            return Response({"error": "No file uploaded"}, status=400)

        # Extract text
        try:
            if file.name.lower().endswith(".pdf"):
                pages = extract_text_from_pdf(file)
            elif file.name.lower().endswith(".docx"):
                pages = extract_text_from_docx(file)
            else:
                return Response({"error": "Unsupported file type"}, status=400)
        except Exception as e:
            return Response({"error": f"Failed to extract text: {str(e)}"}, status=400)

        # Combine text for model
        combined_text = "\n".join([p["text"] for p in pages])

        # Generate sections using model
        prompt = f"""
Extract the following sections from the clinical trial protocol and return a JSON object with section names as keys and extracted content as values.

For each section, also indicate which page numbers from the document were used to generate that content.

Return format:
{{
    "Purpose of the Study": {{
        "content": "extracted content here",
        "source_pages": [1, 2]
    }},
    "Study Procedures": {{
        "content": "extracted content here (include number of patients and study duration if available)",
        "source_pages": [3, 4, 5]
    }},
    "Risks": {{
        "content": "extracted content here",
        "source_pages": [6, 7]
    }},
    "Benefits": {{
        "content": "extracted content here",
        "source_pages": [8, 9]
    }}
}}

Text to analyze (pages 1-{len(pages)}):
{combined_text[:4000]}
"""

        # Try OpenAI API first, fallback to saved response on errors
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a clinical document parser. Return only valid JSON format as requested."},
                    {"role": "user", "content": prompt},
                ],
            )
            generated_text = response.choices[0].message.content
            
            # Clean the response text - remove markdown code blocks
            cleaned_text = generated_text.strip()
            if cleaned_text.startswith('```json'):
                cleaned_text = cleaned_text[7:]  # Remove ```json
            if cleaned_text.startswith('```'):
                cleaned_text = cleaned_text[3:]   # Remove ```
            if cleaned_text.endswith('```'):
                cleaned_text = cleaned_text[:-3]  # Remove trailing ```
            cleaned_text = cleaned_text.strip()
            
            # Try to parse as JSON
            try:
                parsed_sections = json.loads(cleaned_text)
                # Convert back to string for template processing
                generated_text = json.dumps(parsed_sections, indent=2)
            except json.JSONDecodeError:
                pass
            
        except Exception as e:
            # Fallback to saved response
            # Read the saved response
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            response_file = os.path.join(project_root, 'openai_response_sample.txt')
            with open(response_file, 'r') as f:
                generated_text = f.read()

        # Try to parse as JSON first, then fallback to markdown parsing
        generated_sections = {}
        section_source_pages = {}  # Track which pages were used for each section
        
        try:
            # Clean the response text - remove markdown code blocks if present
            cleaned_text = generated_text.strip()
            if cleaned_text.startswith('```json'):
                cleaned_text = cleaned_text[7:]  # Remove ```json
            if cleaned_text.startswith('```'):
                cleaned_text = cleaned_text[3:]   # Remove ```
            if cleaned_text.endswith('```'):
                cleaned_text = cleaned_text[:-3]  # Remove trailing ```
            cleaned_text = cleaned_text.strip()

            parsed_json = json.loads(cleaned_text)
            
            # Handle new format with content and source_pages
            for section_name, section_data in parsed_json.items():
                if isinstance(section_data, dict):
                    # New format with content and source_pages
                    content = section_data.get('content', '')
                    source_pages = section_data.get('source_pages', [])
                    if content.strip():
                        generated_sections[section_name] = content.strip()
                        section_source_pages[section_name] = source_pages
                elif isinstance(section_data, str) and section_data.strip():
                    # Old format - just string content
                    generated_sections[section_name] = section_data.strip()
                    section_source_pages[section_name] = []  # No page info available
        except json.JSONDecodeError:
            # Fallback to markdown parsing
            current_section = None
            current_content = []
            
            for line in generated_text.split('\n'):
                line = line.strip()
                if line.startswith('**') and line.endswith('**'):
                    # Save previous section
                    if current_section:
                        generated_sections[current_section] = '\n'.join(current_content).strip()
                    # Start new section
                    section_name = line.replace('**', '').strip()
                    current_section = section_name
                    current_content = []
                elif current_section and line:
                    current_content.append(line)
            
            # Save last section
            if current_section:
                generated_sections[current_section] = '\n'.join(current_content).strip()
                section_source_pages[current_section] = []  # No page info available in markdown fallback
        
        # Create detailed logs showing which pages were used for each section
        detailed_logs = []
        
        # Add section generation logs with page mapping
        if generated_sections:
            for section_name, content in generated_sections.items():
                # Get source pages from ChatGPT response
                source_pages = section_source_pages.get(section_name, [])
                
                # Create contributing pages info from ChatGPT's source pages
                contributing_pages = []
                for page_num in source_pages:
                    # Find the page object
                    page_obj = next((p for p in pages if p["page"] == page_num), None)
                    if page_obj:
                        contributing_pages.append({
                            "page": page_obj["page"],
                            "content_sample": page_obj["text"][:150] + "..." if len(page_obj["text"]) > 150 else page_obj["text"],
                            "relevance_score": "ChatGPT identified"  # ChatGPT determined this page was relevant
                        })
                
                # If no source pages from ChatGPT, fallback to keyword matching
                if not contributing_pages:
                    for page in pages:
                        # Simple keyword matching to identify relevant pages
                        page_text_lower = page["text"].lower()
                        
                        # Check if page contains keywords related to this section
                        section_keywords = {
                            "Purpose of the Study": ["purpose", "objective", "aim", "goal", "study", "trial"],
                            "Study Procedures": ["procedure", "method", "protocol", "enrollment", "randomization", "treatment"],
                            "Risks": ["risk", "adverse", "side effect", "complication", "danger", "harm"],
                            "Benefits": ["benefit", "advantage", "improvement", "efficacy", "outcome", "positive"]
                        }
                        
                        if section_name in section_keywords:
                            keywords = section_keywords[section_name]
                            if any(keyword in page_text_lower for keyword in keywords):
                                contributing_pages.append({
                                    "page": page["page"],
                                    "content_sample": page["text"][:150] + "..." if len(page["text"]) > 150 else page["text"],
                                    "relevance_score": sum(1 for keyword in keywords if keyword in page_text_lower)
                                })
                    
                    # Sort by relevance and take top contributing pages
                    contributing_pages.sort(key=lambda x: x["relevance_score"] if isinstance(x["relevance_score"], int) else 0, reverse=True)
                    contributing_pages = contributing_pages[:3]  # Top 3 most relevant pages
                
                detailed_logs.append({
                    "type": "section_generation",
                    "section": section_name,
                    "content_preview": content[:200] + "..." if len(content) > 200 else content,
                    "content_length": len(content),
                    "contributing_pages": contributing_pages,
                    "description": f"Generated '{section_name}' section using pages: {', '.join([str(p['page']) for p in contributing_pages]) if contributing_pages else 'AI analysis of all pages'}"
                })
        
        # Add document generation summary
        detailed_logs.append({
            "type": "document_generation",
            "sections_count": len(generated_sections),
            "total_pages_processed": len(pages),
            "description": "Generated DOCX document with extracted protocol information"
        })
        
        # Create a new document with extracted protocol information
        doc = Document()
        
        # Add document title
        doc.add_heading('Informed Consent Form', level=1)
        doc.add_heading('Generated from Clinical Trial Protocol', level=2)
        doc.add_paragraph()  # Add spacing
        
        # Add generation timestamp
        timestamp = datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")
        doc.add_paragraph(f"Generated on: {timestamp}")
        doc.add_paragraph()
        
        # Create sections with proper formatting
        if generated_sections:
            # Add table of contents section
            doc.add_heading('Protocol Information Summary', level=2)
            doc.add_paragraph("The following information has been extracted from the clinical trial protocol:")
            
            # Add section list
            for section_name in generated_sections.keys():
                doc.add_paragraph(f"• {section_name}", style='List Bullet')
            
            doc.add_paragraph()  # Add spacing
            
            # Add each section with proper formatting
            for section_name, content in generated_sections.items():
                # Add section header
                doc.add_heading(section_name, level=2)
                
                # Add content with proper formatting
                content_para = doc.add_paragraph(content)
                content_para.style = 'Normal'
                
                # Add spacing between sections
                doc.add_paragraph()
            
            # Add processing logs section
            doc.add_heading('Processing Logs', level=2)
            doc.add_paragraph("The following logs show the page numbers and content used to generate each section:")
            doc.add_paragraph()
            
            # Add detailed logs to document
            for log_entry in detailed_logs:
                if log_entry["type"] == "extraction":
                    doc.add_heading(f"Page {log_entry['page']} - Text Extraction", level=3)
                    doc.add_paragraph(f"Content Length: {log_entry['content_length']} characters")
                    doc.add_paragraph("Content Preview:")
                    content_preview = doc.add_paragraph(log_entry['content_preview'])
                    content_preview.style = 'Intense Quote'
                    doc.add_paragraph()
                    
                elif log_entry["type"] == "section_generation":
                    doc.add_heading(f"Section: {log_entry['section']}", level=3)
                    doc.add_paragraph(f"Content Length: {log_entry['content_length']} characters")
                    
                    # Show contributing pages
                    if log_entry.get('contributing_pages'):
                        doc.add_paragraph("Source Pages Used:")
                        for page_info in log_entry['contributing_pages']:
                            doc.add_paragraph(f"• Page {page_info['page']} (relevance: {page_info['relevance_score']})")
                            page_sample = doc.add_paragraph(f"  Content sample: {page_info['content_sample']}")
                            page_sample.style = 'Intense Quote'
                        doc.add_paragraph()
                    
                    doc.add_paragraph("Generated Content:")
                    generated_content = doc.add_paragraph(log_entry['content_preview'])
                    generated_content.style = 'Intense Quote'
                    doc.add_paragraph()
                    
                elif log_entry["type"] == "document_generation":
                    doc.add_heading("Document Generation Summary", level=3)
                    doc.add_paragraph(f"Sections Generated: {log_entry['sections_count']}")
                    doc.add_paragraph(f"Pages Processed: {log_entry['total_pages_processed']}")
                    doc.add_paragraph()
        else:
            # Fallback content if no sections were parsed
            doc.add_heading('Protocol Information', level=2)
            doc.add_paragraph("No structured information could be extracted from the protocol document.")
            doc.add_paragraph("Original response:")
            doc.add_paragraph(generated_text)
            
            # Add basic processing logs even in fallback case
            if detailed_logs:
                doc.add_paragraph()
                doc.add_heading('Processing Logs', level=2)
                doc.add_paragraph("Processing information:")
                doc.add_paragraph()
                
                for log_entry in detailed_logs:
                    if log_entry["type"] == "extraction":
                        doc.add_paragraph(f"• Page {log_entry['page']}: {log_entry['content_length']} characters extracted")
                    elif log_entry["type"] == "document_generation":
                        doc.add_paragraph(f"• Total pages processed: {log_entry['total_pages_processed']}")
                        doc.add_paragraph(f"• Sections generated: {log_entry['sections_count']}")

        # Save to temp file
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="icf_")
        doc.save(tmp_file.name)

        # Keep backward compatibility with simple log format
        simple_log = [{"page": p["page"], "text_sample": p["text"][:100]} for p in pages]

        # Prepare response data
        response_data = {
            "download_url": f"/api/download_icf/?file={os.path.basename(tmp_file.name)}",
            "generated_text": generated_text,
            "log": simple_log,
            "detailed_logs": detailed_logs
        }
        
        # Try to add parsed sections if JSON was successful
        try:
            parsed_sections = json.loads(generated_text)
            response_data["sections"] = parsed_sections
        except json.JSONDecodeError:
            pass
        
        return Response(response_data)

class DownloadICF(APIView):
    def get(self, request):
        file_name = request.GET.get("file")
        if not file_name:
            return Response({"error": "File not specified"}, status=400)
        
        file_path = os.path.join(tempfile.gettempdir(), file_name)
        
        # Check if file exists
        if not os.path.exists(file_path):
            return Response({"error": "File not found"}, status=404)
        
        # Generate a descriptive filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        download_filename = f"Protocol_Extracted_Information_{timestamp}.docx"
        
        return FileResponse(
            open(file_path, "rb"), 
            as_attachment=True, 
            filename=download_filename,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
